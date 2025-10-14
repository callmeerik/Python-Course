"""
    Generate an article using open AI
"""
import streamlit as st
from google import genai
from docx import Document
from io import BytesIO

st.set_page_config(
    page_title="Article Generation",
    page_icon="📒",
    layout="wide"
)

st.subheader("Article Generator Using Gemini")

def generate_article(topic):
    try:
        client = genai.Client(api_key=st.secrets.get("GEMINI_API_KEY"))
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=f"""
            You are an expert in writing SEO articles.
            Write an optimized article for SEO about: {topic}, give me only the article text with 
            titles, subtitles, I mean academic style.
            """
        )
    except Exception as ex:
        st.error(f"Error generating article: {ex}")
        return None
    return response.text


# Function to create the word document
def create_word_document(article):
    # instance Document class
    doc = Document()
    # add title to the file
    doc.add_heading("Generated Article", level=0)

    # add paragraph
    doc.add_paragraph(article)

    # create buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# input for topic user
topic = st.text_input("Write the topic for article")

# creat a button fot generate article
if st.button("📋 Generate Article"):
    if topic:
        # create a loader
        with st.spinner("Creating Article..."):
            article = generate_article(topic)
            if article:
                st.success("Article Generated")
                st.markdown("#### Article Preview")
                with st.container(border=True):
                    st.markdown(article)

                # btn for downloading article like word document
                st.download_button(
                    "Download as .docx",
                    data=create_word_document(article),
                    file_name="Article.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )